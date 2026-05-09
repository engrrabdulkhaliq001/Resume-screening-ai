"""
Resume Parser - Step 5
PDF/DOCX resume se information extract karta hai
"""

import re
import json
import spacy
import pdfplumber
from pathlib import Path
from docx import Document

# spacy model load
nlp = spacy.load("en_core_web_sm")

# ===================== SKILLS LIST =====================

SKILLS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "sql",
    "react", "angular", "vue", "nodejs", "django", "flask", "fastapi",
    "machine learning", "deep learning", "nlp", "tensorflow", "pytorch",
    "scikit-learn", "pandas", "numpy", "keras", "opencv",
    "aws", "azure", "gcp", "docker", "kubernetes", "git", "linux",
    "postgresql", "mongodb", "mysql", "redis", "elasticsearch",
    "html", "css", "rest api", "graphql", "microservices",
    "data science", "data analysis", "computer vision",
    "agile", "scrum", "ci/cd", "jenkins", "github", "gitlab",
    "tableau", "power bi", "excel", "spark", "hadoop",
    "php", "ruby", "swift", "kotlin", "flutter", "dart",
    "selenium", "pytest", "junit", "postman",
]

# ===================== TEXT EXTRACTORS =====================

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF error: {e}")
    return text.strip()

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Tables se bhi text lo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"DOCX error: {e}")
    return text.strip()

def extract_text(filepath):
    filepath = Path(filepath)
    if filepath.suffix.lower() == ".pdf":
        return extract_text_from_pdf(filepath)
    elif filepath.suffix.lower() in [".docx", ".doc"]:
        return extract_text_from_docx(filepath)
    elif filepath.suffix.lower() == ".txt":
        return filepath.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: {filepath.suffix}")


# ===================== EXTRACTORS =====================

def extract_email(text):
    pattern = r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
    matches = re.findall(pattern, text)
    return matches[0] if matches else ""

def extract_phone(text):
    patterns = [
        r'(\+92|0092|92)?[-.\s]?3\d{2}[-.\s]?\d{7}',   # Pakistan
        r'(\+1)?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US
        r'(\+\d{1,3})?[-.\s]?\d{10,13}',               # General
    ]
    for pat in patterns:
        match = re.search(pat, text)
        if match:
            return re.sub(r'\s+', '', match.group()).strip()
    return ""

def extract_name(text):
    """Pehli line ya spacy NER se naam nikalo"""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        first_line = lines[0]
        # Agar pehli line mein sirf 2-4 words hain aur email/phone nahi
        words = first_line.split()
        if 2 <= len(words) <= 4 and "@" not in first_line and not any(c.isdigit() for c in first_line):
            return first_line

    # spacy NER fallback
    doc = nlp(text[:500])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return ""

def extract_skills_from_text(text):
    text_lower = text.lower()
    found = []
    for skill in SKILLS:
        if skill in text_lower:
            found.append(skill)
    return found

def extract_experience_years(text):
    patterns = [
        r'(\d+)\+?\s*years?\s*of\s*experience',
        r'(\d+)\+?\s*years?\s*experience',
        r'experience\s*of\s*(\d+)\+?\s*years?',
        r'(\d+)\+?\s*yr[s]?\s*(of\s*)?exp',
    ]
    for pat in patterns:
        match = re.search(pat, text.lower())
        if match:
            return int(match.group(1))
    return None

def extract_education(text):
    degrees = {
        "phd": ["phd", "ph.d", "doctorate", "doctor of"],
        "masters": ["masters", "m.s", "m.sc", "msc", "mba", "m.tech", "ms in"],
        "bachelors": ["bachelors", "b.s", "b.sc", "bsc", "b.tech", "be ", "bs in", "bscs", "bsit", "bsse"],
        "associate": ["associate", "a.s", "a.a"],
        "diploma": ["diploma", "certificate"],
    }
    text_lower = text.lower()
    for degree, keywords in degrees.items():
        for kw in keywords:
            if kw in text_lower:
                return degree
    return "not specified"

def extract_experience_sections(text):
    """Work experience entries nikalo"""
    experiences = []
    lines = text.split("\n")

    # Common job title keywords
    job_keywords = [
        "engineer", "developer", "analyst", "manager", "intern",
        "consultant", "designer", "architect", "lead", "senior",
        "junior", "associate", "officer", "executive", "director",
    ]

    year_pattern = re.compile(r'\b(19|20)\d{2}\b')

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        if any(kw in line_lower for kw in job_keywords) and year_pattern.search(line):
            experiences.append(line.strip())

    return experiences[:5]  # Top 5


# ===================== MAIN PARSER =====================

def parse_resume(filepath):
    """
    Resume file parse karke structured data return karo
    """
    print(f"\nParsing: {filepath}")

    # Text extract karo
    text = extract_text(filepath)
    if not text:
        return {"error": "Could not extract text from file"}

    print(f"  Text extracted: {len(text)} characters")

    # Information extract karo
    result = {
        "file": str(filepath),
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "education": extract_education(text),
        "experience_years": extract_experience_years(text),
        "skills": extract_skills_from_text(text),
        "experience_entries": extract_experience_sections(text),
        "raw_text": text[:1000],  # Preview only
        "full_text": text,
    }

    print(f"  Name     : {result['name']}")
    print(f"  Email    : {result['email']}")
    print(f"  Phone    : {result['phone']}")
    print(f"  Education: {result['education']}")
    print(f"  Exp Years: {result['experience_years']}")
    print(f"  Skills   : {result['skills'][:5]}...")

    return result


# ===================== BATCH PARSER =====================

def parse_all_resumes(folder="resumes"):
    """
    Folder mein saare resumes parse karo
    """
    folder = Path(folder)
    if not folder.exists():
        folder.mkdir()
        print(f"Folder banaya: {folder}/")
        print("Apne resume files yahan rakho aur dobara chalao")
        return []

    supported = [".pdf", ".docx", ".doc", ".txt"]
    files = [f for f in folder.iterdir() if f.suffix.lower() in supported]

    if not files:
        print(f"Koi resume nahi mila {folder}/ mein")
        print("PDF ya DOCX files rakho wahan")
        return []

    print(f"\n{len(files)} resume(s) mile!")
    results = []

    for f in files:
        try:
            parsed = parse_resume(f)
            results.append(parsed)
        except Exception as e:
            print(f"  Error {f.name}: {e}")

    # Save karo
    output = Path("data/parsed_resumes.json")
    save_data = [{k: v for k, v in r.items() if k != "full_text"} for r in results]
    with open(output, "w", encoding="utf-8") as f:
        json.dump(save_data, f, indent=2, ensure_ascii=False)

    print(f"\nParsed {len(results)} resumes saved: {output}")
    return results


# ===================== TEST =====================

if __name__ == "__main__":
    print("=== Resume Parser ===")
    print("\nOption 1: Single resume test karo")
    print("  from resume_parser import parse_resume")
    print("  result = parse_resume('my_resume.pdf')")
    print("\nOption 2: Saare resumes parse karo")
    print("  from resume_parser import parse_all_resumes")
    print("  results = parse_all_resumes('resumes')")
    print("\nSeedha test karna hai toh:")
    print("  'resumes' folder banao, PDF rakho, phir chalao:")
    print("  python -c \"from resume_parser import parse_all_resumes; parse_all_resumes()\"")

    # Auto test if resumes folder exists
    resumes_dir = Path("resumes")
    if resumes_dir.exists() and any(resumes_dir.iterdir()):
        print("\nResumes folder mila! Parsing shuru...")
        parse_all_resumes()